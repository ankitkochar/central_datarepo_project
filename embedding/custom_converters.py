from haystack import component
from bs4 import BeautifulSoup, Comment, NavigableString
from typing import List, Union, Dict, Any, Optional
from haystack import Document
import docx
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
from haystack import Document, component
from .pdf_to_markdown import convert_pdf_to_markdown_using_paddleocr, convert_pdf_to_markdown_using_pytesseract
from markdownify import MarkdownConverter
import re
import io
import os
import boto3
from botocore.client import Config
import gc

from dotenv import load_dotenv

load_dotenv()

aws_access_key = os.environ.get("AWS_ACCESS_KEY")
aws_secret_key = os.environ.get("AWS_SECRET_KEY")
aws_region = os.environ.get("AWS_REGION")
ocr_type = os.environ.get("OCR_TYPE")


def get_s3_file_content(file_path: str) -> Optional[bytes]:
    bucket_name = "cld-data-extraction"
    key = file_path[len("https://cld-data-extraction.s3.amazonaws.com/") :]

    s3_client = boto3.client(
        "s3",
        aws_access_key_id=aws_access_key,
        aws_secret_access_key=aws_secret_key,
        config=Config(signature_version="s3v4"),
        region_name=aws_region,
    )

    try:
        with io.BytesIO() as file_object:
            s3_client.download_fileobj(bucket_name, key, file_object)
            file_object.seek(0)
            return file_object.read()
    except Exception as e:
        print(f"S3 URL: {file_path} Got S3 Error: ", e)
        return None
    finally:
        s3_client.close()


@component
class PDFToDocumentConverter:
    @component.output_types(documents=List[Document])
    def run(
        self,
        sources: List[Union[str, Path]],
        meta: Optional[Union[Dict[str, Any], List[Dict[str, Any]]]] = None,
    ):
        if meta is None:
            meta = {}
        documents = []
        for file_path in sources:
            meta["file_path"] = file_path
            file_content = get_s3_file_content(file_path)
            if file_content:
                try:
                    if ocr_type=="PADDLE":
                        extracted_text = convert_pdf_to_markdown_using_paddleocr(file_content, file_path)
                    else:
                        extracted_text = convert_pdf_to_markdown_using_pytesseract(file_content, file_path)
                    doc = Document(content=extracted_text, meta=meta.copy())
                    documents.append(doc)
                except Exception as e:
                    print(f"S3 URL: {file_path} Got PDF Converter Error: ", e)
                finally:
                    del file_content
                    del extracted_text
                    gc.collect()  # Force garbage collection
        return {"documents": documents}


@component
class URLToDocumentConverterMarkdownify:
    @staticmethod
    def md_soup(soup, **options):
        return MarkdownConverter(**options).convert_soup(soup)

    @component.output_types(documents=List[Document])
    def run(
        self,
        sources: List[Union[str, Path]],
        meta: Optional[Union[Dict[str, Any], List[Dict[str, Any]]]] = None,
    ):
        if meta is None:
            meta = {}
        documents = []
        for file_path in sources:
            meta["file_path"] = file_path
            html_content = get_s3_file_content(file_path)
            if html_content:
                try:
                    html = html_content.decode("utf-8")
                    soup = BeautifulSoup(html, "lxml")

                    for table in soup.find_all("table"):
                        table.insert_before(soup.new_string("[TABLE]"))
                        table.insert_after(soup.new_string("[/TABLE]"))

                    md_content = self.md_soup(
                        soup,
                        strip=["a", "nav", "footer", "img"],
                        heading_style="ATX",
                        autolinks=False,
                        wrap=False,
                        newline_style="\n",
                        escape_asterisks=True,
                        escape_underscores=True,
                    )
                    md_content = re.sub(r"\n{2,}", "\n", md_content)
                    doc = Document(content=md_content, meta=meta.copy())
                    documents.append(doc)
                except Exception as e:
                    print(f"S3 URL: {file_path} Got HTML Converter Error: ", e)
                finally:
                    del html_content
                    del html
                    del soup
                    del md_content
                    gc.collect()  # Force garbage collection
        return {"documents": documents}


@component
class URLToDocumentConverterTabNewline:
    """
    A component to convert HTML files to documents, extracting text and preserving the raw HTML of tables,
    while removing scripts, styles, and other non-essential elements.
    """

    @staticmethod
    def clean_soup(soup):

        [s.decompose() for s in soup("script")]
        # fmt: off
        # fmt: off
        for tag in soup(
            ["script", "style", "link", "meta", "input", "form", "noscript", "iframe", "img", "svg", "button", "aside", "figure", "fielset", "details", "textarea", "fieldset"]
        ):
            tag.decompose()
        # fmt: on
        # fmt: on
        for element in soup.find_all(string=lambda text: isinstance(text, Comment)):
            element.extract()
        for div in soup.find_all("div", style="display:none;"):
            div.decompose()

    @staticmethod
    def remove_unwanted_attributes(element):
        unwanted_attrs = [
            "target",
            "align",
            "nowrap",
            "scrollbar",
            "onclick",
            "resizable",
            "style",
            "title",
            "statusbar",
            "id",
            "class",
            "toolbar",
            "location",
            "datasheetsvalue",
            "datasheetsnumberformat",
            "cellpadding",
            "border",
            "rel",
        ]

        for tag in element.find_all(True):
            for attr in unwanted_attrs:
                if tag.has_attr(attr):
                    del tag[attr]

    @staticmethod
    def format_table(table):
        rows_text = []
        for tr in table.find_all("tr"):
            cols_text = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
            row_text = " | ".join(cols_text)
            rows_text.append(row_text)
        return "[TABLE]\n" + "\n".join(rows_text) + "\n[/TABLE]"

    @staticmethod
    def extract_content(soup):
        contents = []
        processed_tables = set()

        for element in soup.find_all(
            True,
            recursive=False,
        ):
            if element.name == "table" and element not in processed_tables:
                table_text = URLToDocumentConverterTabNewline.format_table(element)
                contents.append(table_text)
                processed_tables.add(element)
            elif element.find_all("table", recursive=True):
                contents.append(
                    URLToDocumentConverterTabNewline.extract_content(element)
                )
            else:
                text = " ".join(element.stripped_strings)
                if text:
                    contents.append(text)

        return " \n ".join(contents)

    @component.output_types(documents=List[Document])
    def run(
        self,
        sources: List[Union[str, Path]],
        meta: Optional[Union[Dict[str, Any], List[Dict[str, Any]]]] = None,
    ):
        if meta is None:
            meta = {}
        documents = []
        for file_path in sources:
            with open(file_path, "r", encoding="utf-8") as file:
                html = file.read()
            if isinstance(meta, dict):
                if isinstance(file_path, Path):
                    meta_file_path = str(file_path)
                else:
                    meta_file_path = file_path
                meta["file_path"] = meta_file_path
            soup = BeautifulSoup(html, "lxml")
            URLToDocumentConverterTabNewline.clean_soup(soup)
            # print(soup.prettify())
            content = URLToDocumentConverterTabNewline.extract_content(soup)
            doc = Document(content=content, meta=meta)
            documents.append(doc)
        return {"documents": documents}


@component
class URLToDocumentConverter:
    """
    A component to convert HTML file to Document with BS4
    """

    @component.output_types(documents=List[Document])
    def run(
        self,
        sources: List[Union[str, Path]],
        meta: Optional[Union[Dict[str, Any], List[Dict[str, Any]]]] = None,
    ):
        if meta is None:
            meta = {}
        documents = []
        for file_path in sources:
            with open(file_path, "r", encoding="utf-8") as file:
                html = file.read()
            if isinstance(meta, dict):
                if isinstance(file_path, Path):
                    meta_file_path = str(file_path)
                else:
                    meta_file_path = file_path
                meta["file_path"] = meta_file_path
            soup = BeautifulSoup(html, "html.parser")

            text = soup.get_text()
            doc = Document(content=text, meta=meta)
            documents.append(doc)
        return {"documents": documents}


@component
class DocxToTextConverter:
    """
    A component to convert docx file to Document
    """

    @component.output_types(documents=List[Document])
    def run(
        self,
        sources: List[Union[str, Path]],
        meta: Optional[Union[Dict[str, Any], List[Dict[str, Any]]]] = None,
    ):
        if meta is None:
            meta = {}
        documents = []
        for file_path in sources:
            source = file_path
            bucket_name = "cld-data-extraction"
            key = source[len("https://cld-data-extraction.s3.amazonaws.com/") :]

            s3_client = boto3.client(
                "s3",
                aws_access_key_id=aws_access_key,
                aws_secret_access_key=aws_secret_key,
                config=Config(signature_version="s3v4"),
                region_name=aws_region,
            )

            try:
                s3_object = s3_client.get_object(Bucket=bucket_name, Key=key)
                file_content = s3_object["Body"].read()
                doc = docx.Document(io.BytesIO(file_content))

            except Exception as e:
                print("Got error:- ", e)

            text = ""
            for para in doc.paragraphs:
                text += para.text
            doc = Document(content=text, meta=meta)
            documents.append(doc)
        return {"documents": documents}
